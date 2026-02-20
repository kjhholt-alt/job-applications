"""
Resume & Cover Letter DOCX Generator
Generates crisp, professional Word documents.
Save as PDF from Word for perfect quality.
Usage: python generate_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, bottom={"sz": 6, "val": "single"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        for key, val in attrs.items():
            element.set(qn(f'w:{key}'), str(val))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_horizontal_line(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_narrow_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_header(doc, name, address, phone_email):
    """Centered name and contact."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(address)
    run2.font.size = Pt(10.5)
    run2.font.name = "Calibri"

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(2)
    run3 = p3.add_run(phone_email)
    run3.font.size = Pt(10.5)
    run3.font.name = "Calibri"


def add_section_header(doc, title):
    """Bold section header with underline."""
    add_horizontal_line(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"


def add_job_header(doc, title, dates):
    """Bold title left, italic dates right using a table."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.2)

    # Remove table borders
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'right', 'bottom']:
            el = OxmlElement(f'w:{edge}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            tcBorders.append(el)
        tcPr.append(tcBorders)

    # Left cell - title
    left = table.rows[0].cells[0].paragraphs[0]
    left.paragraph_format.space_before = Pt(6)
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Right cell - dates
    right = table.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_before = Pt(6)
    right.paragraph_format.space_after = Pt(0)
    run2 = right.add_run(dates)
    run2.italic = True
    run2.font.size = Pt(11)
    run2.font.name = "Calibri"

    # Remove table spacing
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    spacing = OxmlElement('w:tblCellSpacing')
    spacing.set(qn('w:w'), '0')
    spacing.set(qn('w:type'), 'dxa')
    tblPr.append(spacing)


def add_location(doc, location):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(location)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.35)
    # Clear default text and add our own
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"


def add_project_bullet(doc, name, description):
    """Bullet with bold project name + regular description."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.35)
    p.clear()
    run_name = p.add_run(name + " ")
    run_name.bold = True
    run_name.font.size = Pt(10.5)
    run_name.font.name = "Calibri"
    run_desc = p.add_run(description)
    run_desc.font.size = Pt(10.5)
    run_desc.font.name = "Calibri"


def add_skills_line(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    run_label.font.size = Pt(10.5)
    run_label.font.name = "Calibri"
    run_text = p.add_run(text)
    run_text.font.size = Pt(10.5)
    run_text.font.name = "Calibri"


def add_subsection_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def add_italic_intro(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"


def add_body_text(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


# =====================================================
# RESUME 1: Finance Analytics & AI Manager
# =====================================================
def generate_finance_ai_resume():
    doc = Document()
    set_narrow_margins(doc)

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)

    add_header(doc, "Kruz J. Holt", "23 S Jacob Drive, Parkview, Iowa 52748", "563-343-1255  |  kjh.holt@gmail.com")

    # Education
    add_section_header(doc, "Education")
    add_job_header(doc, "University of Iowa, Iowa City, IA", "December 2020")
    add_body_text(doc, "B.A. Finance, Tippie College of Business (Direct Admit)")
    add_body_text(doc, "College of Tippie GPA: 3.5/4.0")

    # Professional Experience
    add_section_header(doc, "Professional Experience")

    add_job_header(doc, "John Deere Seeding \u2013 Operations Accounting Analyst", "February 2025 \u2013 Present")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "Own factory overhead budgets and lead monthly variance analysis reviews with plant leadership, translating operational changes into financial forecasts and actionable insights.")
    add_bullet(doc, "Develop automated reporting solutions using Power BI and Advanced Excel, driving cross-functional process improvement and reducing manual reporting effort.")
    add_bullet(doc, "Perform month-end close activities including expense classification, variance analysis, and standards alignment for accurate financial reporting.")
    add_bullet(doc, "Track and model factory headcount changes, connecting workforce data to overhead forecasts and supporting labor-related financial planning.")
    add_bullet(doc, "Support capital project tracking (AFE/AFL), validating requests and monitoring spend against approved budgets.")

    add_job_header(doc, "John Deere Harvester \u2013 New Product Cost Analyst", "2024 \u2013 2025")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "Designed and deployed a Power BI dashboard used by engineering and program management to analyze monthly cost trends across new product programs.")
    add_bullet(doc, "Extracted and analyzed Bill of Materials (BoM) data from SAP for combines, front-end equipment, and seeding programs, providing cost estimates for new product development.")
    add_bullet(doc, "Tracked capital expenditures across multiple programs, ensuring accurate financial management and budget compliance.")

    add_job_header(doc, "John Deere Financial \u2013 Accounts Receivable Analyst", "2023 \u2013 2024")
    add_location(doc, "Des Moines, Iowa")
    add_bullet(doc, "Pioneered automation of financial workflows using Power Automate and SQL, eliminating manual steps and delivering measurable efficiency gains across the AR function.")
    add_bullet(doc, "Managed completion of 57 monthly financial tasks with zero errors, demonstrating precision in high-volume financial operations.")
    add_bullet(doc, "Developed training programs for new team members, establishing standardized processes and fostering a knowledge-sharing culture.")
    add_bullet(doc, "Performed reconciliations in SAP and uploaded certified data in Blackline; leveraged ESSBASE to analyze daily and monthly lease volumes.")

    add_job_header(doc, "John Deere Corporate \u2013 General Ledger / Fixed Asset Shared Services", "Summer 2022 \u2013 2023")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "Conducted reconciliations for all international units, ensuring accurate financial reporting across global operations.")
    add_bullet(doc, "Analyzed leases and collaborated with international units on calculations and reconciliation; compiled data for annual 10-K report.")

    add_job_header(doc, "Principal Global Investors \u2013 Investment Operations Intern", "Summer 2018 \u2013 Winter 2019")
    add_location(doc, "Des Moines, Iowa")
    add_bullet(doc, "Calculated performance of hundreds of composites with millions in AUM for GIPS compliance using Bloomberg, Excel, and Sylvan.")
    add_bullet(doc, "Collaborated across Fixed Income, Equity, and Real Estate boutiques to produce verified performance data.")

    # AI & Technology Projects
    add_section_header(doc, "AI & Technology Projects")
    add_italic_intro(doc, "Independently designed, built, and deployed production applications using generative AI and modern web technologies.")
    add_project_bullet(doc, "AI Finance Brief \u2014", "Next.js SaaS generating personalized financial briefings via Claude AI API, with authentication (NextAuth) and email delivery (Resend). Deployed on Vercel.")
    add_project_bullet(doc, "Trade Journal \u2014", "Full-stack trade analytics platform (Next.js + Python FastAPI) with AI-powered analysis, performance visualization, and automated insight generation.")
    add_project_bullet(doc, "Automated Trading System \u2014", "Python quantitative trading bot with real-time data pipeline, scoring engine (Kelly criterion, VWAP, momentum), risk management, and Supabase analytics. Deployed on Railway.")
    add_project_bullet(doc, "CRM Platforms \u2014", "React + Django full-stack CRMs for outdoor services and franchise management with dashboards, PDF generation, and offline-first mobile design.")

    # Leadership
    add_section_header(doc, "Leadership & AI Advocacy")
    add_job_header(doc, "AI Lunch & Learn Facilitator \u2013 John Deere Ag & Turf Accounting", "2024 \u2013 Present")
    add_bullet(doc, "Lead quarterly AI sessions for the entire Ag & Turf accounting function, covering generative AI developments, finance-specific use cases, and live tool demonstrations to drive enterprise AI adoption.")
    add_subsection_title(doc, "Development Committee Member \u2013 John Deere Accounting Organization")
    add_bullet(doc, "Contribute to talent development and organizational growth initiatives.")

    # Skills
    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "AI & Analytics", "Generative AI (Claude API, prompt engineering), Power BI, Power Automate, Advanced Excel, SQL, Python (pandas), quantitative modeling")
    add_skills_line(doc, "Finance Systems", "SAP, Blackline, ESSBASE, Bloomberg")
    add_skills_line(doc, "Development", "Python, TypeScript/JavaScript, React, Next.js, Django, FastAPI, REST APIs, Supabase (PostgreSQL), Git")
    add_skills_line(doc, "Cloud & DevOps", "Vercel, Railway, CI/CD, agile development practices")

    out = "C:/Users/Kruz/Desktop/Kruz_Holt_Resume_Deloitte_Finance_AI.docx"
    doc.save(out)
    print(f"Created: {out}")


# =====================================================
# RESUME 2: AI Solutions Leader
# =====================================================
def generate_ai_solutions_resume():
    doc = Document()
    set_narrow_margins(doc)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)

    add_header(doc, "Kruz J. Holt", "23 S Jacob Drive, Parkview, Iowa 52748", "563-343-1255  |  kjh.holt@gmail.com")

    # Education
    add_section_header(doc, "Education")
    add_job_header(doc, "University of Iowa, Iowa City, IA", "December 2020")
    add_body_text(doc, "B.A. Finance, Tippie College of Business (Direct Admit)")
    add_body_text(doc, "College of Tippie GPA: 3.5/4.0")

    # AI Portfolio (LEADS)
    add_section_header(doc, "AI & Technology Portfolio")
    add_italic_intro(doc, "Independently designed, architected, and deployed 10+ production applications using generative AI, full-stack web technologies, and cloud infrastructure. Hands-on across the entire lifecycle: requirements, architecture, engineering, deployment, and optimization.")

    add_subsection_title(doc, "AI-Powered SaaS Applications (2024 \u2013 Present)")
    add_project_bullet(doc, "AI Finance Brief \u2014", "Production SaaS: Next.js 14, Claude AI API, NextAuth, Resend. Personalized financial briefings via generative AI with structured output parsing. Deployed on Vercel with CI/CD.")
    add_project_bullet(doc, "AI Chess Coach \u2014", "Production app: Next.js 14, chess.js, Lichess API, Claude AI. Real-time game analysis with AI coaching. Custom prompt architecture for domain-specific reasoning.")
    add_project_bullet(doc, "Trade Journal \u2014", "Full-stack: Next.js 14 + Python FastAPI. AI-powered trade analysis, performance visualization (Recharts), automated insights. Vercel + Railway deployment.")

    add_subsection_title(doc, "Automated Trading System \u2014 \"MoneyPrinter\" (2025 \u2013 Present)")
    add_bullet(doc, "Architected and deployed a Python quantitative trading bot for prediction markets with real-time data pipelines, multi-strategy execution, and automated risk management.")
    add_bullet(doc, "Built scoring engine incorporating momentum analysis, VWAP signals, and Kelly criterion position sizing with dynamic confidence thresholds per asset class.")
    add_bullet(doc, "Engineered parallel strategy execution (4.1x speedup), batch order processing, and per-asset rate limiters \u2014 reduced cycle times from 17s to 2.6s.")
    add_bullet(doc, "Implemented production monitoring: Supabase analytics backend, automated watchdog with health checks and self-healing deployment, Discord bot for remote management.")

    add_subsection_title(doc, "Full-Stack Platforms & Infrastructure (2024 \u2013 Present)")
    add_project_bullet(doc, "CRM Systems \u2014", "React 19 + Django full-stack CRMs with lead discovery engines, job scheduling, PDF invoicing, offline-first mobile (IndexedDB), and analytics dashboards.")
    add_project_bullet(doc, "Admin Dashboard \u2014", "Next.js 15 agent management system with Supabase backend, real-time monitoring, and Zustand state management.")
    add_project_bullet(doc, "Automation \u2014", "Browser automation (Puppeteer, Playwright), web scraping pipelines (Selenium, BeautifulSoup), Discord-integrated AI agent system with cron scheduling.")

    # Professional Experience
    add_section_header(doc, "Professional Experience")

    add_job_header(doc, "John Deere Seeding \u2013 Operations Accounting Analyst", "February 2025 \u2013 Present")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "Own factory overhead budgets and deliver monthly variance analysis to plant leadership, translating operational data into financial forecasts and actionable insights.")
    add_bullet(doc, "Build automated reporting dashboards using Power BI and Advanced Excel, supporting cross-functional process improvement and data-driven decision making.")

    add_job_header(doc, "John Deere Harvester \u2013 New Product Cost Analyst", "2024 \u2013 2025")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "Designed and deployed Power BI dashboard for engineering and program management, enabling self-service monthly cost analysis across new product programs.")
    add_bullet(doc, "Extracted and analyzed SAP BoM data, providing cost estimates for new product development across combines, front-end equipment, and seeding.")

    add_job_header(doc, "John Deere Financial \u2013 Accounts Receivable Analyst", "2023 \u2013 2024")
    add_location(doc, "Des Moines, Iowa")
    add_bullet(doc, "Pioneered automation of financial workflows using Power Automate and SQL, delivering measurable efficiency gains. Managed 57 monthly tasks; SAP reconciliations and ESSBASE analysis.")

    add_job_header(doc, "John Deere Corporate \u2013 GL / Fixed Asset Shared Services", "Summer 2022 \u2013 2023")
    add_location(doc, "Davenport, Iowa")
    add_bullet(doc, "International unit reconciliations, lease analysis, and financial data compilation for annual 10-K SEC reporting.")

    add_job_header(doc, "Principal Global Investors \u2013 Investment Operations Intern", "Summer 2018 \u2013 Winter 2019")
    add_location(doc, "Des Moines, Iowa")
    add_bullet(doc, "Performance analysis across hundreds of composites using Bloomberg, Excel, and Sylvan for GIPS compliance.")

    # Leadership
    add_section_header(doc, "Leadership & AI Advocacy")
    add_job_header(doc, "AI Lunch & Learn Facilitator \u2013 John Deere Ag & Turf Accounting", "2024 \u2013 Present")
    add_bullet(doc, "Lead quarterly AI sessions for the Ag & Turf accounting organization, demonstrating GenAI tools, use cases, and driving enterprise AI adoption.")

    # Skills
    add_section_header(doc, "Technical Skills")
    add_skills_line(doc, "AI/ML & GenAI", "Claude AI API, prompt engineering & architecture, quantitative modeling, AI workflow automation, generative AI application development")
    add_skills_line(doc, "Languages", "Python, TypeScript/JavaScript, SQL, HTML/CSS")
    add_skills_line(doc, "Frontend", "React, Next.js (14/15), Tailwind CSS, shadcn/ui, Framer Motion, Recharts")
    add_skills_line(doc, "Backend", "Django (DRF), FastAPI, Node.js, REST API design & implementation")
    add_skills_line(doc, "Data & Cloud", "Supabase (PostgreSQL), pandas, Vercel, Railway, Git/GitHub, CI/CD, Docker")
    add_skills_line(doc, "Finance Tools", "SAP, Blackline, ESSBASE, Bloomberg, Power BI, Power Automate, Advanced Excel")

    out = "C:/Users/Kruz/Desktop/Kruz_Holt_Resume_Deloitte_AI_Solutions.docx"
    doc.save(out)
    print(f"Created: {out}")


# =====================================================
# COVER LETTERS
# =====================================================
def add_cover_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"


def generate_finance_ai_cover():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Kruz J. Holt")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    for line in ["23 S Jacob Drive, Parkview, Iowa 52748", "563-343-1255  |  kjh.holt@gmail.com"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"

    doc.add_paragraph()  # spacer

    for line in ["February 20, 2026", "", "Deloitte Recruiting", "Finance Transformation \u2014 Finance Analytics & AI Manager", "", "Dear Hiring Team,"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    paras = [
        "I\u2019m applying for the Finance Analytics & AI Manager position on Deloitte\u2019s Finance Transformation team. What drew me to this role is that it sits at the exact intersection of my career: I\u2019ve spent four years in core finance at John Deere \u2014 FP&A, reporting, cost analysis, month-end close \u2014 and I\u2019ve independently built and deployed production AI applications that solve real business problems. I don\u2019t just talk about AI in finance. I build it.",
        "At John Deere, I\u2019ve worked across the finance function: overhead budgeting and variance analysis at the factory level, new product cost estimation using SAP BoM data, accounts receivable automation with Power Automate and SQL, and international reconciliations for 10-K reporting. I designed Power BI dashboards that replaced manual reporting cycles and are used daily by engineers and program managers. This gives me the finance process fluency your clients need \u2014 I understand the pain points because I\u2019ve lived them.",
        "Outside of my day job, I\u2019ve architected and deployed over 10 full-stack applications using generative AI. My AI Finance Brief app uses Claude\u2019s API to generate personalized financial briefings \u2014 the kind of AI-enabled finance insight your team delivers to CFOs. My automated trading system processes real-time market data through a quantitative scoring engine with dynamic risk management, deployed on cloud infrastructure with automated monitoring. These aren\u2019t side projects \u2014 they\u2019re production systems with real users and real money. I\u2019ve worked across the full AI solution lifecycle: identifying the use case, designing the architecture, building the solution, deploying to cloud, and iterating based on performance data.",
        "I lead quarterly AI Lunch & Learn sessions for John Deere\u2019s entire Ag & Turf accounting function, covering generative AI news, practical finance use cases, and live tool demonstrations. Helping others understand and adopt AI is something I genuinely enjoy \u2014 and it\u2019s directly relevant to the client-facing and mentorship aspects of this role.",
        "I recognize my background is non-traditional for this level. My finance degree isn\u2019t STEM-designated, and my years of experience are weighted toward the early side of your range. But a candidate who has both deep finance process knowledge AND the ability to personally build the AI solutions \u2014 not just manage their delivery \u2014 brings unusual value to a Finance Analytics & AI team.",
        "I\u2019d welcome the opportunity to discuss how my combination of finance expertise and hands-on AI development can contribute to Deloitte\u2019s Finance Transformation practice.",
    ]

    for text in paras:
        add_cover_paragraph(doc, text)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Kruz J. Holt")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    out = "C:/Users/Kruz/Desktop/Kruz_Holt_Cover_Letter_Deloitte_Finance_AI.docx"
    doc.save(out)
    print(f"Created: {out}")


def generate_ai_solutions_cover():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Kruz J. Holt")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"

    for line in ["23 S Jacob Drive, Parkview, Iowa 52748", "563-343-1255  |  kjh.holt@gmail.com"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"

    doc.add_paragraph()

    for line in ["February 20, 2026", "", "Deloitte Recruiting", "Human Capital \u2014 AI Solutions Leader", "", "Dear Hiring Team,"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Calibri"

    doc.add_paragraph()

    paras = [
        "I\u2019m applying for the AI Solutions Leader position within Deloitte\u2019s Human Capital practice. While my years of experience don\u2019t match a traditional senior hire, I want to make the case that what I bring is harder to find: I\u2019m a finance professional who independently architects, builds, and deploys production AI systems end-to-end \u2014 and I understand business problems well enough to know which ones AI should actually solve.",
        "Over the past two years, I\u2019ve designed, built, and deployed 10+ production applications spanning the complete technology stack. My AI-powered SaaS applications use generative AI APIs with custom prompt architecture, modern frontend frameworks (Next.js, React, TypeScript), Python backends (Django, FastAPI), and cloud deployment (Vercel, Railway, Supabase). My automated trading system is the most technically demanding: a Python quantitative engine with real-time data pipelines, parallel strategy execution (4.1x speedup), dynamic risk management using Kelly criterion, and production monitoring with automated health checks and self-healing deployment. I\u2019ve managed the full lifecycle on every project \u2014 requirements, architecture, engineering, deployment, and continuous optimization.",
        "Building production systems teaches you things that theoretical knowledge doesn\u2019t. I\u2019ve designed prompt architectures for domain-specific AI reasoning, built data pipelines that handle structured and unstructured sources, implemented rate limiting and caching for API-heavy applications, and made infrastructure decisions around cost, latency, and reliability. I\u2019ve also learned when NOT to use AI \u2014 my trading system removed a signal that was circular because I measured its actual impact. This kind of practical AI judgment is what your clients need.",
        "Four years at John Deere gave me deep understanding of enterprise finance: FP&A, cost analysis, financial reporting, SEC compliance, and process automation. I\u2019ve automated workflows with Power Automate and SQL, built Power BI dashboards for cross-functional teams, and worked in SAP across multiple business units. When consulting with leaders on AI solutions, I can speak their language \u2014 I\u2019ve sat in their seat.",
        "I lead quarterly AI Lunch & Learn sessions for John Deere\u2019s Ag & Turf accounting organization, demonstrating generative AI tools and driving enterprise adoption. Making technical concepts accessible to non-technical stakeholders is something I do today and would bring to client engagements at Deloitte.",
        "I\u2019m transparent that my experience profile is earlier-career than your posting describes. But the AI landscape has shifted \u2014 practical building experience with generative AI matters more than years on a resume, and candidates who can both sell an AI vision and personally build the solution are rare. I\u2019m one of them.",
        "I\u2019d be excited to discuss how my hands-on AI development experience and finance domain knowledge can contribute to Deloitte\u2019s Human Capital AI practice.",
    ]

    for text in paras:
        add_cover_paragraph(doc, text)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    run = p.add_run("Kruz J. Holt")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    out = "C:/Users/Kruz/Desktop/Kruz_Holt_Cover_Letter_Deloitte_AI_Solutions.docx"
    doc.save(out)
    print(f"Created: {out}")


if __name__ == "__main__":
    generate_finance_ai_resume()
    generate_ai_solutions_resume()
    generate_finance_ai_cover()
    generate_ai_solutions_cover()
    print("\nDone! All 4 DOCX files on your Desktop.")
    print("Open in Word and Save As PDF for crisp, high-quality output.")
